from __future__ import annotations
import base64
import io
import json
import os
import re
from collections import deque
from pathlib import Path
from typing import Literal
from uuid import uuid4

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
import logfire
from opentelemetry import trace
from pydantic import BaseModel, ConfigDict, Field, model_validator
from pypdf import PdfReader
from pydantic_ai import Agent
from pydantic_ai.agent import AgentRunResult
from pydantic_ai.models.anthropic import AnthropicModel
import requests

from urllib import error as urlerror
from urllib import request as urlrequest

load_dotenv()

logfire_key = os.getenv("LOGFIRE_KEY")
if logfire_key and not os.getenv("LOGFIRE_TOKEN"):
    os.environ["LOGFIRE_TOKEN"] = logfire_key

logfire.configure()

BASE_DIR = Path(__file__).resolve().parents[1]
STATIC_DIR = BASE_DIR / "static"


def _parse_env_list(value: str) -> list[str]:
    return [item.strip() for item in value.split(",") if item.strip()]


def _resolve_frontend_dir() -> Path:
    env_dir = os.getenv("FRONTEND_DIST_DIR", "").strip()
    candidates: list[Path] = []
    if env_dir:
        candidates.append(Path(env_dir))
    candidates.extend(
        [
            BASE_DIR / "frontend" / "dist",
            BASE_DIR / "dist",
            STATIC_DIR,
        ]
    )
    for candidate in candidates:
        if candidate.is_dir() and (candidate / "index.html").is_file():
            return candidate
    return STATIC_DIR


FRONTEND_DIR = _resolve_frontend_dir()
FRONTEND_DIR_RESOLVED = FRONTEND_DIR.resolve()
FRONTEND_ASSETS_DIR = FRONTEND_DIR / "assets"
FRONTEND_DEV_SERVER_URL = os.getenv("FRONTEND_DEV_SERVER_URL", "").strip()

CORS_ORIGINS = _parse_env_list(
    os.getenv("CORS_ORIGINS", os.getenv("VITE_DEV_ORIGINS", ""))
)
if not CORS_ORIGINS:
    CORS_ORIGINS = ["http://localhost:5173", "http://127.0.0.1:5173"]
BYPASS_OUTPUT_VALIDATION = True
MODEL_MAX_TOKENS = int(os.getenv("MODEL_MAX_TOKENS", "12000"))
MAX_PROMPT_TOKENS = int(os.getenv("MAX_PROMPT_TOKENS", str(MODEL_MAX_TOKENS)))
MODEL_SETTINGS = {"max_tokens": MODEL_MAX_TOKENS}
LOCAL_AI_URL = os.getenv("LOCAL_AI_URL")
LOCAL_AI_MODEL = os.getenv("LOCAL_AI_MODEL", "dictalm-he:latest")
LOCAL_AI_TIMEOUT = float(os.getenv("LOCAL_AI_TIMEOUT", "180"))
LOCAL_AI_NUM_PREDICT = int(os.getenv("LOCAL_AI_NUM_PREDICT", str(MODEL_MAX_TOKENS)))
LOCAL_AI_NUM_CTX = int(os.getenv("LOCAL_AI_NUM_CTX", str(MODEL_MAX_TOKENS)))
LOCAL_AI_NEMO_URL = os.getenv("LOCAL_AI_NEMO_URL", "http://10.19.0.70:8000")
LOCAL_AI_NEMO_MODEL = os.getenv(
    "LOCAL_AI_NEMO_MODEL",
    "mistralai/Mistral-Nemo-Instruct-FP8-2407",
)
LOCAL_AI_NEMO_TIMEOUT = float(os.getenv("LOCAL_AI_NEMO_TIMEOUT", "60"))
LOCAL_AI_NEMO_TEMPERATURE = float(os.getenv("LOCAL_AI_NEMO_TEMPERATURE", "0.2"))

SYSTEM_PROMPT = (
    "You are an automated grading assistant.\n"
    "You will receive two raw text inputs: (1) a guide that contains questions and grading expectations, and (2) a student's answer text.\n"
    "The guide may include guidance documents, rubrics, and raw questions (with or without teacher answers); use all of it to infer questions and scoring expectations.\n"
    "The texts may have inconsistent or missing formatting; do not rely on numbering or delimiters. Infer question boundaries and answer alignment from meaning and context.\n"
    "Return ONLY a JSON object with exactly seven keys: student_name, questions, answers, grades, comments, highlights, criteria.\n"
    "student_name must be the student's name if it is explicitly present in the student answers (or guide); otherwise use an empty string.\n"
    "Do not invent a name.\n"
    "questions must be an object keyed by sequential question numbers as strings (\"1\", \"2\", ...) with the full question text.\n"
    "If a question contains subparts (e.g., A/B, א/ב, Paragraph A/B, סעיף א/ב), keep them in the question text as separate lines.\n"
    "Use the same labels/language as the guide and prefix each subpart with its label (e.g., \"Subquestion A: ...\" or \"סעיף א: ...\").\n"
    "answers must be an object keyed by the same question numbers with the aligned student answer text.\n"
    "grades must be an object keyed by the same question numbers; each value is an object with numeric score and max_score based on the guide's point breakdown (use 10 if unspecified).\n"
    "comments must be an object keyed by the same question numbers with a short grading explanation (1-2 sentences).\n"
    "highlights must be an object keyed by the same question numbers; each value is a list of highlight objects.\n"
    "Each highlight object must include: start (0-based index, inclusive), end (0-based index, exclusive), effect (\"add\" or \"deduct\"), reason (short phrase).\n"
    "For \"deduct\" highlights, include points (positive number). For \"add\" highlights, omit points or set it to 0.\n"
    "The start/end indices must refer to the exact substring in answers[k].\n"
    "Scoring rule: for each question, start from max_score and subtract ONLY the sum of points for \"deduct\" highlights.\n"
    "\"add\" highlights are positive/strengths and MUST NOT increase the score.\n"
    "For every question where score < max_score, include deduct highlights whose points sum exactly to (max_score - score).\n"
    "Do not use empty ranges; always select a short exact span from the answer when deducting points.\n"
    "If an element is missing from the answer, still deduct points by highlighting the closest relevant sentence and explain the missing element in the reason.\n"
    "Also include \"add\" highlights for strengths with a reason (no points needed). Do NOT omit positive highlights.\n"
    "criteria must be a list of rubric criteria that are NOT actual questions (e.g., overall quality, organization, language, bibliography).\n"
    "Each criteria item must include: title, score, max_score, comment, and optional items (subcriteria) with the same fields.\n"
    "If the guide contains rubric bullets like \"Overall quality (10 points)\" with sub-items, treat them as criteria, not questions.\n"
    "Keep the question order as it appears in the guide. If an answer is missing, output an empty string, score 0, and a short comment.\n"
    "Do not output any extra keys, metadata, markdown, or explanation - only JSON.\n"
    "Ensure the JSON is well-formed and valid.\n"
    "Don't write anything outside the JSON object.\n"
    "The first character of your response must be '{' and the last character must be '}'.\n"
    f"Make it less than {MAX_PROMPT_TOKENS} tokens in total."
)

USER_PROMPT_TEMPLATE = (
    "GUIDE (raw text):\n"
    "{guide}\n\n"
    "STUDENT ANSWERS (raw text):\n"
    "{student_answers}"
)


class GradeRequest(BaseModel):
    guide: str = ""
    student_answers: str = ""
    guide_pdf_base64: str = ""
    guide_pdf_base64_list: list[str] = Field(default_factory=list)
    guide_docx_base64_list: list[str] = Field(default_factory=list)
    student_pdf_base64: str = ""
    student_docx_base64: str = ""
    custom_prompt: str = ""


class GradeDetail(BaseModel):
    model_config = ConfigDict(extra="forbid")

    score: float
    max_score: float

    @model_validator(mode="after")
    def validate_scores(self) -> "GradeDetail":
        if self.max_score <= 0:
            raise ValueError("max_score must be positive")
        if not 0 <= self.score <= self.max_score:
            raise ValueError("score must be between 0 and max_score")
        return self


class Highlight(BaseModel):
    model_config = ConfigDict(extra="forbid")

    start: int
    end: int
    effect: Literal["add", "deduct"]
    points: float | None = None
    reason: str

    @model_validator(mode="after")
    def validate_highlight(self) -> "Highlight":
        if self.start < 0 or self.end <= self.start:
            raise ValueError("Invalid highlight range")
        if not str(self.reason).strip():
            raise ValueError("Highlight reason is required")
        points = 0.0 if self.points is None else float(self.points)
        if self.effect == "deduct":
            if points <= 0:
                raise ValueError("Deduct highlight points must be positive")
        elif points < 0:
            raise ValueError("Highlight points must be non-negative")
        self.points = points
        return self


class CriterionItem(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str
    score: float | None = None
    max_score: float | None = None
    comment: str = ""

    @model_validator(mode="after")
    def validate_criterion_item(self) -> "CriterionItem":
        if not str(self.title).strip():
            raise ValueError("Criterion title is required")
        if self.max_score is not None and self.max_score <= 0:
            raise ValueError("Criterion max_score must be positive when provided")
        if self.score is not None and self.max_score is not None:
            if not 0 <= self.score <= self.max_score:
                raise ValueError("Criterion score must be between 0 and max_score")
        return self


class CriterionGroup(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str
    score: float | None = None
    max_score: float | None = None
    comment: str = ""
    items: list[CriterionItem] = Field(default_factory=list)

    @model_validator(mode="after")
    def validate_criterion_group(self) -> "CriterionGroup":
        if not str(self.title).strip():
            raise ValueError("Criterion group title is required")
        if self.max_score is not None and self.max_score <= 0:
            raise ValueError("Criterion group max_score must be positive when provided")
        if self.score is not None and self.max_score is not None:
            if not 0 <= self.score <= self.max_score:
                raise ValueError("Criterion group score must be between 0 and max_score")
        return self


class GradingView(BaseModel):
    model_config = ConfigDict(extra="forbid")

    student_name: str = ""
    questions: dict[str, str] = Field(default_factory=dict)
    answers: dict[str, str] = Field(default_factory=dict)
    grades: dict[str, GradeDetail] = Field(default_factory=dict)
    comments: dict[str, str] = Field(default_factory=dict)
    highlights: dict[str, list[Highlight]] = Field(default_factory=dict)
    criteria: list[CriterionGroup] = Field(default_factory=list)

    @model_validator(mode="after")
    def validate_payload(self) -> "GradingView":
        normalized_questions: dict[str, str] = {}
        normalized_answers: dict[str, str] = {}
        normalized_grades: dict[str, GradeDetail] = {}
        normalized_comments: dict[str, str] = {}
        normalized_highlights: dict[str, list[Highlight]] = {}
        student_name = str(self.student_name or "").strip()
        if student_name:
            student_name = re.sub(
                r"^(name|student name|student)[:\s]+", "", student_name, flags=re.IGNORECASE
            ).strip()
        if student_name.lower() in {"unknown", "n/a", "na", "none", "not provided"}:
            student_name = ""
        if student_name:
            words = student_name.split()
            if len(words) > 8:
                student_name = " ".join(words[:8])

        for key, value in self.questions.items():
            key_str = str(key)
            if re.fullmatch(r"[1-9]\d*", key_str) is None:
                raise ValueError(f"Invalid question key: {key_str}")
            normalized_questions[key_str] = str(value).strip()

        for key, value in self.answers.items():
            key_str = str(key)
            if re.fullmatch(r"[1-9]\d*", key_str) is None:
                raise ValueError(f"Invalid answer key: {key_str}")
            normalized_answers[key_str] = str(value)

        for key, value in self.grades.items():
            key_str = str(key)
            if re.fullmatch(r"[1-9]\d*", key_str) is None:
                raise ValueError(f"Invalid grade key: {key_str}")
            normalized_grades[key_str] = value

        for key, value in self.comments.items():
            key_str = str(key)
            if re.fullmatch(r"[1-9]\d*", key_str) is None:
                raise ValueError(f"Invalid comment key: {key_str}")
            normalized_comments[key_str] = str(value).strip()

        for key, value in self.highlights.items():
            key_str = str(key)
            if re.fullmatch(r"[1-9]\d*", key_str) is None:
                raise ValueError(f"Invalid highlight key: {key_str}")
            normalized_highlights[key_str] = value

        if not normalized_grades:
            raise ValueError("grades is required")

        key_sets = [
            set(normalized_questions.keys()),
            set(normalized_answers.keys()),
            set(normalized_grades.keys()),
            set(normalized_comments.keys()),
        ]
        if any(not key_set for key_set in key_sets):
            raise ValueError("questions, answers, grades, and comments must be populated")

        if len({frozenset(key_set) for key_set in key_sets}) != 1:
            raise ValueError("questions, answers, grades, and comments keys must match")

        keys = sorted(int(key) for key in normalized_grades.keys())
        if keys != list(range(1, len(keys) + 1)):
            raise ValueError("Question keys must be consecutive starting at 1")

        for key in normalized_grades.keys():
            normalized_highlights.setdefault(key, [])
            answer_text = normalized_answers.get(key, "")
            valid_highlights: list[Highlight] = []
            for highlight in normalized_highlights[key]:
                if highlight.end <= len(answer_text):
                    valid_highlights.append(highlight)
            normalized_highlights[key] = valid_highlights

        for key, grade in normalized_grades.items():
            answer_text = normalized_answers.get(key, "")
            max_score = grade.max_score
            if not str(answer_text).strip():
                score = 0.0
            else:
                deduct_total = sum(
                    highlight.points
                    for highlight in normalized_highlights.get(key, [])
                    if highlight.effect == "deduct" and highlight.points > 0
                )
                score = max(0.0, max_score - deduct_total)
            normalized_grades[key] = GradeDetail(score=score, max_score=max_score)

        normalized_criteria: list[CriterionGroup] = []
        for criterion in self.criteria:
            title = str(criterion.title).strip()
            if not title:
                continue
            items: list[CriterionItem] = []
            for item in criterion.items:
                item_title = str(item.title).strip()
                if not item_title:
                    continue
                score = item.score
                max_score = item.max_score
                if max_score is not None and max_score > 0 and score is not None:
                    score = max(0.0, min(score, max_score))
                elif score is not None:
                    score = max(0.0, score)
                items.append(
                    CriterionItem(
                        title=item_title,
                        score=score,
                        max_score=max_score,
                        comment=str(item.comment or "").strip(),
                    )
                )

            score = criterion.score
            max_score = criterion.max_score
            if max_score is not None and max_score > 0 and score is not None:
                score = max(0.0, min(score, max_score))
            elif score is not None:
                score = max(0.0, score)

            normalized_criteria.append(
                CriterionGroup(
                    title=title,
                    score=score,
                    max_score=max_score,
                    comment=str(criterion.comment or "").strip(),
                    items=items,
                )
            )

        self.questions = normalized_questions
        self.answers = normalized_answers
        self.grades = normalized_grades
        self.comments = normalized_comments
        self.highlights = normalized_highlights
        self.criteria = normalized_criteria
        self.student_name = student_name
        return self


class RawOutputResponse(BaseModel):
    raw_output: str


_agent: Agent | None = None
_raw_outputs: dict[str, str] = {}
_raw_output_order: deque[str] = deque()
_raw_output_limit = 20


def _strip_code_fences(text: str) -> str:
    stripped = text.strip()
    if not stripped.startswith("```"):
        return stripped

    if stripped.endswith("```") and "\n" not in stripped:
        inner = stripped[3:-3].strip()
        if inner.lower().startswith("json"):
            inner = inner[4:].lstrip()
        return inner

    lines = stripped.splitlines()
    if lines and lines[0].lstrip().startswith("```"):
        lines = lines[1:]
    if lines and lines[-1].strip().startswith("```"):
        lines = lines[:-1]
    if lines and lines[0].strip().lower() == "json":
        lines = lines[1:]
    return "\n".join(lines).strip()


def _strip_data_url(value: str) -> str:
    if not value:
        return ""
    text = value.strip()
    if text.startswith("data:"):
        comma_index = text.find(",")
        if comma_index != -1:
            return text[comma_index + 1 :]
    return text


def _extract_pdf_text(pdf_data: str, *, label: str) -> str:
    if not pdf_data:
        return ""
    try:
        payload = _strip_data_url(pdf_data)
        if not payload:
            return ""
        pdf_bytes = base64.b64decode(payload, validate=False)
    except Exception as exc:
        logfire.warning("Failed to decode PDF payload", label=label, error=str(exc))
        return ""
    if not pdf_bytes:
        return ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        pages: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)
        return "\n\n".join(pages).strip()
    except Exception as exc:
        logfire.warning("Failed to extract PDF text", label=label, error=str(exc))
        return ""


def _extract_docx_text(docx_data: str, *, label: str) -> str:
    if not docx_data:
        return ""
    try:
        payload = _strip_data_url(docx_data)
        if not payload:
            return ""
        docx_bytes = base64.b64decode(payload, validate=False)
    except Exception as exc:
        logfire.warning("Failed to decode DOCX payload", label=label, error=str(exc))
        return ""
    if not docx_bytes:
        return ""
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        logfire.warning("DOCX extraction unavailable (python-docx not installed)", label=label, error=str(exc))
        return ""
    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        logfire.warning("Failed to read DOCX document", label=label, error=str(exc))
        return ""
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def _resolve_grade_inputs(payload: GradeRequest) -> tuple[str, str]:
    guide_text = str(payload.guide or "")
    student_text = str(payload.student_answers or "")
    pdf_chunks: list[str] = []
    if payload.guide_pdf_base64:
        pdf_chunks.append(_extract_pdf_text(payload.guide_pdf_base64, label="guide"))
    if payload.guide_pdf_base64_list:
        for index, item in enumerate(payload.guide_pdf_base64_list, start=1):
            if not item:
                continue
            pdf_chunks.append(_extract_pdf_text(item, label=f"guide_{index}"))
    if pdf_chunks:
        combined_pdf = "\n\n".join(chunk for chunk in pdf_chunks if chunk.strip())
        if combined_pdf:
            if guide_text.strip():
                guide_text = f"{guide_text}\n\n{combined_pdf}"
            else:
                guide_text = combined_pdf
    docx_chunks: list[str] = []
    if payload.guide_docx_base64_list:
        for index, item in enumerate(payload.guide_docx_base64_list, start=1):
            if not item:
                continue
            docx_chunks.append(_extract_docx_text(item, label=f"guide_docx_{index}"))
    if docx_chunks:
        combined_docx = "\n\n".join(chunk for chunk in docx_chunks if chunk.strip())
        if combined_docx:
            if guide_text.strip():
                guide_text = f"{guide_text}\n\n{combined_docx}"
            else:
                guide_text = combined_docx
    if not student_text.strip() and payload.student_pdf_base64:
        student_text = _extract_pdf_text(payload.student_pdf_base64, label="student")
    if not student_text.strip() and payload.student_docx_base64:
        student_text = _extract_docx_text(payload.student_docx_base64, label="student_docx")
    return guide_text, student_text


def _ndjson_event(payload: dict[str, object]) -> bytes:
    return (json.dumps(payload, ensure_ascii=False) + "\n").encode("utf-8")



def local_ai(
    prompt: str,
    url: str | None = None,
    timeout: float = LOCAL_AI_TIMEOUT,
    model: str | None = None,
) -> str:
    target = url or LOCAL_AI_URL or "http://127.0.0.1:11434/api/generate"
    if not target:
        raise ValueError("local_ai requires a url or LOCAL_AI_URL")

    payload_data: dict[str, object] = {"prompt": prompt, "stream": False}
    model_name = model or LOCAL_AI_MODEL
    if model_name:
        payload_data["model"] = model_name
    options: dict[str, object] = {}
    if LOCAL_AI_NUM_PREDICT > 0:
        options["num_predict"] = LOCAL_AI_NUM_PREDICT
    if LOCAL_AI_NUM_CTX > 0:
        options["num_ctx"] = LOCAL_AI_NUM_CTX
    if options:
        payload_data["options"] = options
        


    payload = json.dumps(payload_data).encode("utf-8")
    request = urlrequest.Request(
        target,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urlrequest.urlopen(request, timeout=timeout) as response:
            body = response.read().decode("utf-8", errors="replace")
    except urlerror.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Local AI request failed ({exc.code}): {detail}") from exc
    except urlerror.URLError as exc:
        raise RuntimeError(f"Local AI request failed: {exc.reason}") from exc

    try:
        parsed = json.loads(body)
    except json.JSONDecodeError:
        lines = [line for line in body.splitlines() if line.strip()]
        if lines:
            chunks: list[str] = []
            for line in lines:
                try:
                    item = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if isinstance(item, dict):
                    chunk = item.get("response") or item.get("content")
                    if isinstance(chunk, str):
                        chunks.append(chunk)
            if chunks:
                return "".join(chunks)
        return body

    if isinstance(parsed, dict):
        response_text = parsed.get("response")
        if isinstance(response_text, str):
            return response_text
        content_text = parsed.get("content")
        if isinstance(content_text, str):
            return content_text
        choices = parsed.get("choices")
        if isinstance(choices, list) and choices:
            first = choices[0]
            if isinstance(first, dict):
                message = first.get("message")
                if isinstance(message, dict) and isinstance(message.get("content"), str):
                    return message["content"]
                if isinstance(first.get("text"), str):
                    return first["text"]

    return body





# (tu as déjà ces constantes)
# LOCAL_AI_NEMO_URL = "http://10.19.0.70:8000"   # idéalement SANS /v1
# LOCAL_AI_NEMO_MODEL = "mistralai/Mistral-Nemo-Instruct-FP8-2407"
# LOCAL_AI_NEMO_TIMEOUT = 60
# LOCAL_AI_NEMO_TEMPERATURE = 0.2

def local_ai_nemo(
    prompt: str | None = None,
    messages: list[dict[str, str]] | None = None,
    base: str | None = None,
    timeout: float = LOCAL_AI_NEMO_TIMEOUT,
    model: str | None = None,
    temperature: float = LOCAL_AI_NEMO_TEMPERATURE,
    max_tokens: int = 1024,          # ✅ clé pour éviter les valeurs négatives
    top_p: float | None = None,
) -> str:
    target_base = (base or LOCAL_AI_NEMO_URL or "").rstrip("/")
    if not target_base:
        raise ValueError("local_ai_nemo requires a base url or LOCAL_AI_NEMO_URL")

    # ✅ si on te donne déjà .../v1, on évite /v1/v1
    if target_base.endswith("/v1"):
        target_base = target_base[:-3]

    if messages is None:
        if prompt is None:
            raise ValueError("local_ai_nemo requires prompt or messages")
        messages = [{"role": "user", "content": prompt}]

    # ✅ sécurité: max_tokens doit être > 0
    if not isinstance(max_tokens, int) or max_tokens <= 0:
        max_tokens = 256  # fallback sûr

    payload_data: dict[str, object] = {
        "messages": messages,
        "temperature": float(temperature),
        "max_tokens": max_tokens,     # ✅ on l’envoie explicitement
        "stream": False,
    }

    if top_p is not None:
        payload_data["top_p"] = float(top_p)

    model_name = model or LOCAL_AI_NEMO_MODEL
    if model_name:
        payload_data["model"] = model_name

    target = f"{target_base}/v1/chat/completions"

    print("POST", target)
    print(json.dumps(payload_data, ensure_ascii=False, indent=2))

    try:
        response = requests.post(target, json=payload_data, timeout=timeout)
        response.raise_for_status()
    except requests.HTTPError as exc:
        # response existe ici
        print("REQUEST BODY:", getattr(response.request, "body", None))
        print("REQUEST HEADERS:", getattr(response.request, "headers", None))
        raise RuntimeError(
            f"Local AI Nemo request failed ({response.status_code}): {response.text}"
        ) from exc
    except requests.RequestException as exc:
        raise RuntimeError(f"Local AI Nemo request failed: {exc}") from exc

    try:
        parsed = response.json()
    except ValueError:
        return response.text

    if isinstance(parsed, dict):
        choices = parsed.get("choices")
        if isinstance(choices, list) and choices:
            first = choices[0]
            if isinstance(first, dict):
                message = first.get("message")
                if isinstance(message, dict) and isinstance(message.get("content"), str):
                    return message["content"]
                if isinstance(first.get("text"), str):
                    return first["text"]

        error = parsed.get("error")
        if isinstance(error, dict) and isinstance(error.get("message"), str):
            raise RuntimeError(f"Local AI Nemo request failed: {error['message']}")

    return response.text




def _coerce_float(value: object) -> float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        match = re.search(r"-?\d+(?:\.\d+)?", value)
        if match:
            return float(match.group(0))
        return None
    return None


def _extract_json_object(text: str) -> dict | None:
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
        return None
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        try:
            candidate = text[start : end + 1]
            parsed = json.loads(candidate)
            return parsed if isinstance(parsed, dict) else None
        except json.JSONDecodeError:
            candidate = _sanitize_json_text(text[start : end + 1])
            candidate = re.sub(r",\s*([}\]])", r"\1", candidate)
            try:
                parsed = json.loads(candidate)
                return parsed if isinstance(parsed, dict) else None
            except json.JSONDecodeError:
                return None


def _sanitize_json_text(text: str) -> str:
    sanitized: list[str] = []
    in_string = False
    escape = False
    for char in text:
        if in_string:
            if escape:
                escape = False
                sanitized.append(char)
                continue
            if char == "\\":
                escape = True
                sanitized.append(char)
                continue
            if char == "\"":
                in_string = False
                sanitized.append(char)
                continue
            if char == "\n":
                sanitized.append("\\n")
                continue
            if char == "\r":
                sanitized.append("\\r")
                continue
            if char == "\t":
                sanitized.append("\\t")
                continue
            if ord(char) < 0x20:
                sanitized.append(f"\\u{ord(char):04x}")
                continue
            sanitized.append(char)
            continue
        if char == "\"":
            in_string = True
        sanitized.append(char)
    return "".join(sanitized)


def _limit_words(text: str, max_words: int, fallback: str) -> str:
    words = text.strip().split()
    if not words:
        return fallback
    return " ".join(words[:max_words])


def _limit_comment_words(comment: str) -> str:
    return _limit_words(comment, 24, "No feedback.")


def _coerce_comment(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        for key in ("comment", "feedback", "note", "text"):
            if key in value:
                return str(value[key])
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _coerce_student_name(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        for key in ("student_name", "student", "name"):
            if key in value:
                value = value[key]
                break
    text = str(value).strip()
    if not text:
        return ""
    text = re.sub(r"^(name|student name|student)[:\s]+", "", text, flags=re.IGNORECASE).strip()
    lowered = text.lower()
    if lowered in {"unknown", "n/a", "na", "none", "not provided"}:
        return ""
    words = text.split()
    if len(words) > 8:
        text = " ".join(words[:8])
    return text


def _coerce_score_pair(value: object) -> tuple[float, float] | None:
    if isinstance(value, dict):
        score = _coerce_float(value.get("score"))
        max_score = _coerce_float(value.get("max_score")) or _coerce_float(value.get("max"))
        if score is None:
            score = _coerce_float(value.get("grade"))
        if score is None:
            return None
        if max_score is None:
            max_score = 10.0
        return score, max_score
    if isinstance(value, str):
        match = re.search(r"(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?)", value)
        if match:
            return float(match.group(1)), float(match.group(2))
    score = _coerce_float(value)
    if score is None:
        return None
    return score, 10.0


def _coerce_questions_map(source: object) -> dict[str, str]:
    if not isinstance(source, dict):
        return {}
    return {str(key): str(value).strip() for key, value in source.items()}


def _coerce_answers_map(source: object) -> dict[str, str]:
    if not isinstance(source, dict):
        return {}
    return {str(key): str(value) for key, value in source.items()}


def _coerce_grades_map(source: object) -> dict[str, dict[str, float]]:
    if not isinstance(source, dict):
        return {}

    grades: dict[str, dict[str, float]] = {}
    for key, value in source.items():
        pair = _coerce_score_pair(value)
        if pair is None:
            continue
        score, max_score = pair
        if max_score <= 0:
            continue
        score = max(0.0, min(score, max_score))
        grades[str(key)] = {"score": score, "max_score": max_score}

    return grades


def _coerce_comments_map(source: object) -> dict[str, str]:
    if not isinstance(source, dict):
        return {}

    comments: dict[str, str] = {}
    for key, value in source.items():
        comments[str(key)] = _limit_comment_words(_coerce_comment(value))

    return comments


def _coerce_highlight(value: object) -> dict[str, object] | None:
    if not isinstance(value, dict):
        return None
    start = _coerce_float(value.get("start"))
    end = _coerce_float(value.get("end"))
    if start is None or end is None:
        return None
    effect = str(value.get("effect") or "").lower()
    if effect in ("added", "add", "plus", "positive"):
        effect = "add"
    elif effect in ("deducted", "deduct", "minus", "negative"):
        effect = "deduct"
    else:
        return None
    raw_points = _coerce_float(value.get("points"))
    if effect == "deduct":
        if raw_points is None or raw_points == 0:
            return None
        points = abs(raw_points)
    else:
        points = abs(raw_points) if raw_points is not None and raw_points > 0 else 0.0
    if end < start:
        return None
    if end == start:
        return None
    reason = str(value.get("reason") or value.get("note") or "").strip()
    if not reason:
        return None
    return {
        "start": int(start),
        "end": int(end),
        "effect": effect,
        "points": abs(points),
        "reason": reason,
    }


def _coerce_highlights_map(source: object) -> dict[str, list[dict[str, object]]]:
    if not isinstance(source, dict):
        return {}

    highlights: dict[str, list[dict[str, object]]] = {}
    for key, value in source.items():
        if isinstance(value, list):
            cleaned: list[dict[str, object]] = []
            for item in value:
                highlight = _coerce_highlight(item)
                if highlight is not None:
                    cleaned.append(highlight)
            highlights[str(key)] = cleaned
        else:
            highlight = _coerce_highlight(value)
            if highlight is not None:
                highlights[str(key)] = [highlight]

    return highlights


def _coerce_criterion_item(value: object) -> dict[str, object] | None:
    if not isinstance(value, dict):
        return None
    title = str(
        value.get("title")
        or value.get("name")
        or value.get("label")
        or value.get("criterion")
        or value.get("item")
        or ""
    ).strip()
    if not title:
        return None
    score = _coerce_float(value.get("score") or value.get("grade") or value.get("points"))
    max_score = _coerce_float(value.get("max_score") or value.get("max") or value.get("total") or value.get("points_total"))
    if max_score is not None and max_score <= 0:
        max_score = None
    if score is not None and max_score is not None:
        score = max(0.0, min(score, max_score))
    elif score is not None:
        score = max(0.0, score)
    comment = _limit_comment_words(
        _coerce_comment(value.get("comment") or value.get("feedback") or value.get("note") or value.get("reason"))
    ).strip()
    return {
        "title": title,
        "score": score,
        "max_score": max_score,
        "comment": comment,
    }


def _coerce_criterion_group(value: object) -> dict[str, object] | None:
    if not isinstance(value, dict):
        return None
    title = str(
        value.get("title")
        or value.get("name")
        or value.get("label")
        or value.get("criterion")
        or value.get("group")
        or ""
    ).strip()
    if not title:
        return None

    items_source = (
        value.get("items")
        or value.get("subcriteria")
        or value.get("sub_criteria")
        or value.get("criteria")
        or value.get("parts")
        or []
    )
    items: list[dict[str, object]] = []
    if isinstance(items_source, list):
        for item in items_source:
            coerced = _coerce_criterion_item(item)
            if coerced is not None:
                items.append(coerced)
    elif isinstance(items_source, dict):
        for key, item in items_source.items():
            if isinstance(item, dict):
                item.setdefault("title", key)
                coerced = _coerce_criterion_item(item)
            else:
                coerced = _coerce_criterion_item({"title": key, "score": item})
            if coerced is not None:
                items.append(coerced)

    score = _coerce_float(value.get("score") or value.get("grade") or value.get("points"))
    max_score = _coerce_float(value.get("max_score") or value.get("max") or value.get("total") or value.get("points_total"))
    if max_score is None and items:
        totals = [item.get("max_score") for item in items if isinstance(item.get("max_score"), (int, float))]
        max_score = float(sum(totals)) if totals else None
    if score is None and items:
        totals = [item.get("score") for item in items if isinstance(item.get("score"), (int, float))]
        score = float(sum(totals)) if totals else None
    if max_score is not None and max_score <= 0:
        max_score = None
    if score is not None and max_score is not None:
        score = max(0.0, min(score, max_score))
    elif score is not None:
        score = max(0.0, score)

    comment = _limit_comment_words(
        _coerce_comment(value.get("comment") or value.get("feedback") or value.get("note") or value.get("reason"))
    ).strip()

    return {
        "title": title,
        "score": score,
        "max_score": max_score,
        "comment": comment,
        "items": items,
    }


def _coerce_criteria_list(source: object) -> list[dict[str, object]]:
    if source is None:
        return []
    criteria: list[dict[str, object]] = []
    if isinstance(source, list):
        for entry in source:
            coerced = _coerce_criterion_group(entry)
            if coerced is not None:
                criteria.append(coerced)
    elif isinstance(source, dict):
        for key, entry in source.items():
            if isinstance(entry, dict):
                entry.setdefault("title", key)
                coerced = _coerce_criterion_group(entry)
            else:
                coerced = _coerce_criterion_group({"title": key, "score": entry})
            if coerced is not None:
                criteria.append(coerced)
    return criteria


_SUBQUESTION_KEY_RE = re.compile(r"^(\d+)\s*([A-Za-zא-ת])$", re.IGNORECASE)


def _normalize_sub_label(label: str) -> str:
    if not label:
        return ""
    return label.upper() if label.isascii() else label


def _split_question_key(key: str) -> tuple[str | None, str | None]:
    key = str(key or "").strip()
    if re.fullmatch(r"[1-9]\d*", key):
        return key, ""
    match = _SUBQUESTION_KEY_RE.fullmatch(key)
    if match:
        return match.group(1), _normalize_sub_label(match.group(2))
    return None, None


def _has_label_prefix(text: str, label: str) -> bool:
    if not text or not label:
        return False
    label_escaped = re.escape(label)
    patterns = [
        rf"^\s*{label_escaped}\s*[:\).\-]",
        rf"^\s*(Paragraph|Part|Section|Subquestion|Sub-question|Sub question)\s*{label_escaped}\s*[:\).\-]",
        rf"^\s*(פסקה|סעיף|חלק|תת-שאלה)\s*{label_escaped}\s*[:\).\-]",
    ]
    return any(re.match(pattern, text, re.IGNORECASE) for pattern in patterns)


def _merge_sub_answers_with_highlights(
    subs: dict[str, dict[str, object]],
    sub_order: list[str],
) -> tuple[str, list[dict[str, object]]]:
    separator = "\n\n"
    parts: list[str] = []
    merged: list[dict[str, object]] = []
    offset = 0
    for index, label in enumerate(sub_order):
        sub = subs.get(label, {})
        text = str(sub.get("answers", ""))
        prefix = f"{label}: "
        if index > 0:
            offset += len(separator)
        piece = prefix + text
        parts.append(piece)
        for highlight in sub.get("highlights", []) or []:
            try:
                start = int(highlight.get("start", -1))
                end = int(highlight.get("end", -1))
            except (TypeError, ValueError, AttributeError):
                continue
            merged.append(
                {
                    **highlight,
                    "start": start + offset + len(prefix),
                    "end": end + offset + len(prefix),
                }
            )
        offset += len(piece)
    return separator.join(parts), merged


def _merge_subquestion_payload(
    questions: dict[str, str],
    answers: dict[str, str],
    grades: dict[str, dict[str, float]],
    comments: dict[str, str],
    highlights: dict[str, list[dict[str, object]]],
) -> tuple[
    dict[str, str],
    dict[str, str],
    dict[str, dict[str, float]],
    dict[str, str],
    dict[str, list[dict[str, object]]],
]:
    buckets: dict[str, dict[str, object]] = {}

    def assign(source: dict[str, object], field: str) -> None:
        for key, value in source.items():
            base, label = _split_question_key(key)
            if base is None:
                continue
            bucket = buckets.setdefault(base, {"main": {}, "subs": {}, "sub_order": []})
            if label:
                sub_order: list[str] = bucket["sub_order"]
                if label not in sub_order:
                    sub_order.append(label)
                subs: dict[str, dict[str, object]] = bucket["subs"]
                subs.setdefault(label, {})[field] = value
            else:
                main: dict[str, object] = bucket["main"]
                main[field] = value

    assign(questions, "questions")
    assign(answers, "answers")
    assign(grades, "grades")
    assign(comments, "comments")
    assign(highlights, "highlights")

    merged_questions: dict[str, str] = {}
    merged_answers: dict[str, str] = {}
    merged_grades: dict[str, dict[str, float]] = {}
    merged_comments: dict[str, str] = {}
    merged_highlights: dict[str, list[dict[str, object]]] = {}

    for base_key in sorted(buckets.keys(), key=lambda key: int(key)):
        bucket = buckets[base_key]
        main: dict[str, object] = bucket["main"]
        subs: dict[str, dict[str, object]] = bucket["subs"]
        sub_order: list[str] = bucket["sub_order"]

        question_text = str(main.get("questions", "")).strip()
        sub_lines: list[str] = []
        for label in sub_order:
            sub_text = str(subs.get(label, {}).get("questions", "")).strip()
            if not sub_text:
                continue
            if _has_label_prefix(sub_text, label):
                sub_lines.append(sub_text)
            else:
                sub_lines.append(f"{label}: {sub_text}")
        if sub_lines:
            question_text = (question_text + "\n" if question_text else "") + "\n".join(sub_lines)
        merged_questions[base_key] = question_text

        use_sub_answers = any("answers" in subs.get(label, {}) for label in sub_order)
        if use_sub_answers:
            answer_text, merged_sub_highlights = _merge_sub_answers_with_highlights(subs, sub_order)
            merged_answers[base_key] = answer_text
            merged_highlights[base_key] = merged_sub_highlights
        else:
            merged_answers[base_key] = str(main.get("answers", ""))
            merged_highlights[base_key] = list(main.get("highlights", []) or [])

        use_sub_comments = any("comments" in subs.get(label, {}) for label in sub_order)
        if use_sub_comments:
            comment_parts: list[str] = []
            for label in sub_order:
                comment_text = str(subs.get(label, {}).get("comments", "")).strip()
                if not comment_text:
                    continue
                if _has_label_prefix(comment_text, label):
                    comment_parts.append(comment_text)
                else:
                    comment_parts.append(f"{label}: {comment_text}")
            merged_comments[base_key] = "\n".join(comment_parts) if comment_parts else str(
                main.get("comments", "")
            ).strip()
        else:
            merged_comments[base_key] = str(main.get("comments", "")).strip()

        sub_grades = [
            subs.get(label, {}).get("grades")
            for label in sub_order
            if "grades" in subs.get(label, {})
        ]
        if sub_grades:
            total_score = sum(
                value.get("score", 0.0) for value in sub_grades if isinstance(value, dict)
            )
            total_max = sum(
                value.get("max_score", 0.0) for value in sub_grades if isinstance(value, dict)
            )
            if total_max > 0:
                total_score = max(0.0, min(total_score, total_max))
                merged_grades[base_key] = {"score": total_score, "max_score": total_max}
            elif "grades" in main:
                merged_grades[base_key] = main["grades"]
        elif "grades" in main:
            merged_grades[base_key] = main["grades"]

    return merged_questions, merged_answers, merged_grades, merged_comments, merged_highlights


def _extract_grading_view(raw_output: str) -> GradingView | None:
    data = _extract_json_object(raw_output)
    if not isinstance(data, dict):
        return None

    questions_source = data.get("questions") or data.get("separated_questions") or data.get("titles")
    answers_source = data.get("answers") or data.get("student_answers")
    grades_source = data.get("grades") or data.get("grades_for_each_question")
    comments_source = data.get("comments") or data.get("comments_for_each_question")
    highlights_source = data.get("highlights") or {}
    student_source = (
        data.get("student_name")
        or data.get("student")
        or data.get("studentName")
        or data.get("name")
    )
    criteria_source = (
        data.get("criteria")
        or data.get("rubric")
        or data.get("rubric_criteria")
        or data.get("rubric_items")
        or data.get("criteria_groups")
    )

    questions = _coerce_questions_map(questions_source)
    answers = _coerce_answers_map(answers_source)
    grades = _coerce_grades_map(grades_source)
    comments = _coerce_comments_map(comments_source)
    highlights = _coerce_highlights_map(highlights_source)
    criteria = _coerce_criteria_list(criteria_source)
    student_name = _coerce_student_name(student_source)

    if not grades:
        return None

    questions, answers, grades, comments, highlights = _merge_subquestion_payload(
        questions,
        answers,
        grades,
        comments,
        highlights,
    )

    if not grades:
        return None

    sorted_keys = sorted(grades.keys(), key=lambda k: int(k))
    for key in sorted_keys:
        questions.setdefault(key, f"Question {key}")
        answers.setdefault(key, "")
        comments.setdefault(key, "No feedback.")
        highlights.setdefault(key, [])

    questions = {key: questions[key] for key in sorted_keys}
    answers = {key: answers[key] for key in sorted_keys}
    comments = {key: comments[key] for key in sorted_keys}
    highlights = {key: highlights.get(key, []) for key in sorted_keys}

    numeric_keys = [int(key) for key in sorted_keys]
    if numeric_keys != list(range(1, len(sorted_keys) + 1)):
        reindexed = {old: str(index + 1) for index, old in enumerate(sorted_keys)}
        questions = {reindexed[key]: value for key, value in questions.items()}
        answers = {reindexed[key]: value for key, value in answers.items()}
        comments = {reindexed[key]: value for key, value in comments.items()}
        highlights = {reindexed[key]: value for key, value in highlights.items()}
        grades = {reindexed[key]: value for key, value in grades.items()}

    sanitized_highlights: dict[str, list[dict[str, object]]] = {}
    for key, answer_text in answers.items():
        valid: list[dict[str, object]] = []
        for highlight in highlights.get(key, []):
            try:
                start = int(highlight.get("start", -1))
                end = int(highlight.get("end", -1))
            except (TypeError, ValueError):
                continue
            if 0 <= start < end <= len(answer_text):
                reason = str(highlight.get("reason", "")).strip()
                if not reason:
                    continue
                raw_effect = str(highlight.get("effect", "")).lower()
                if raw_effect in ("added", "add", "plus", "positive"):
                    effect = "add"
                elif raw_effect in ("deducted", "deduct", "minus", "negative", "remove"):
                    effect = "deduct"
                else:
                    continue
                raw_points = _coerce_float(highlight.get("points"))
                if effect == "deduct":
                    if raw_points is None or raw_points == 0:
                        continue
                    points = abs(raw_points)
                else:
                    points = abs(raw_points) if raw_points is not None and raw_points > 0 else 0.0
                valid.append(
                    {
                        "start": start,
                        "end": end,
                        "effect": effect,
                        "points": points,
                        "reason": reason,
                    }
                )
        sanitized_highlights[key] = valid

    payload = {
        "student_name": student_name,
        "questions": questions,
        "answers": answers,
        "grades": grades,
        "comments": comments,
        "highlights": sanitized_highlights,
        "criteria": criteria,
    }

    return GradingView.model_validate(payload)


def _extract_raw_output(run_result: AgentRunResult) -> str:
    try:
        response = run_result.response
    except Exception:
        return ""

    output_tool_name = getattr(run_result, "_output_tool_name", None)
    tool_calls = response.tool_calls
    if tool_calls:
        for tool_call in tool_calls:
            if output_tool_name is None or tool_call.tool_name == output_tool_name:
                try:
                    return tool_call.args_as_json_str()
                except Exception:
                    if isinstance(tool_call.args, str):
                        return tool_call.args
                    if tool_call.args is None:
                        return ""
                    return json.dumps(tool_call.args)

        fallback_call = tool_calls[-1]
        try:
            return fallback_call.args_as_json_str()
        except Exception:
            if isinstance(fallback_call.args, str):
                return fallback_call.args
            if fallback_call.args is None:
                return ""
            return json.dumps(fallback_call.args)

    return response.text or ""


def _attach_raw_output(raw_output: str) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.raw_output", raw_output)
        span.set_attribute("grady.raw_output_len", len(raw_output))


def _to_jsonable(value: object) -> object:
    if isinstance(value, BaseModel):
        return value.model_dump()
    if isinstance(value, dict):
        return {str(key): _to_jsonable(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_to_jsonable(item) for item in value]
    return value


def _attach_grades_map(grades: dict[str, object]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.grades", json.dumps(_to_jsonable(grades), ensure_ascii=False))


def _attach_comments_map(comments: dict[str, str]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.comments", json.dumps(comments, ensure_ascii=False))


def _attach_questions_map(questions: dict[str, str]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.questions", json.dumps(questions, ensure_ascii=False))


def _attach_answers_map(answers: dict[str, str]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.answers", json.dumps(answers, ensure_ascii=False))


def _attach_student_name(student_name: str) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.student_name", student_name or "")


def _attach_highlights_map(highlights: dict[str, object]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.highlights", json.dumps(_to_jsonable(highlights), ensure_ascii=False))


def _attach_criteria_list(criteria: list[CriterionGroup]) -> None:
    span = trace.get_current_span()
    if span is not None and span.is_recording():
        span.set_attribute("grady.criteria", json.dumps(_to_jsonable(criteria), ensure_ascii=False))


def _store_raw_output(raw_output: str) -> str:
    run_id = uuid4().hex
    _raw_outputs[run_id] = raw_output
    _raw_output_order.append(run_id)
    if len(_raw_output_order) > _raw_output_limit:
        stale_id = _raw_output_order.popleft()
        _raw_outputs.pop(stale_id, None)
    return run_id


    
    

def get_agent() -> Agent:
    global _agent
    if _agent is not None:
        return _agent

    model = AnthropicModel("claude-sonnet-4-5", provider="gateway")
    _agent = Agent(
        model,
        output_type=GradingView,
        system_prompt=SYSTEM_PROMPT,
        retries=1,
        model_settings=MODEL_SETTINGS,
    )
    return _agent


def _frontend_redirect(request: Request, path: str) -> RedirectResponse | None:
    if not FRONTEND_DEV_SERVER_URL:
        return None
    base = FRONTEND_DEV_SERVER_URL.rstrip("/")
    target = f"{base}/{path.lstrip('/')}" if path else base
    if request.url.query:
        target = f"{target}?{request.url.query}"
    return RedirectResponse(url=target)


def _safe_frontend_path(path: str) -> Path | None:
    if not path:
        return None
    candidate = (FRONTEND_DIR / path).resolve()
    try:
        candidate.relative_to(FRONTEND_DIR_RESOLVED)
    except ValueError:
        return None
    if candidate.is_file():
        return candidate
    return None


def _serve_frontend(request: Request, path: str) -> Response:
    redirect = _frontend_redirect(request, path)
    if redirect is not None:
        return redirect
    file_path = _safe_frontend_path(path)
    if file_path is not None:
        return FileResponse(file_path)
    index_html = (FRONTEND_DIR / "index.html").read_text(encoding="utf-8")
    return HTMLResponse(content=index_html)


app = FastAPI(title="GRADY")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
if hasattr(logfire, "instrument_fastapi"):
    try:
        logfire.instrument_fastapi(app)
    except Exception as exc:
        logfire.warning("FastAPI instrumentation disabled", error=str(exc))

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
if FRONTEND_ASSETS_DIR.is_dir():
    app.mount("/assets", StaticFiles(directory=FRONTEND_ASSETS_DIR), name="assets")


@app.get("/", response_class=HTMLResponse)
def index(request: Request) -> Response:
    return _serve_frontend(request, "")


@app.get("/grade/raw", response_model=RawOutputResponse)
def grade_raw(run_id: str) -> RawOutputResponse:
    raw_output = _raw_outputs.get(run_id)
    if raw_output is None:
        raise HTTPException(status_code=404, detail="Raw output not found")
    return RawOutputResponse(raw_output=raw_output)


@app.post("/grade", response_model=GradingView)
def grade(payload: GradeRequest, response: Response) -> GradingView | Response:
    guide_text, student_text = _resolve_grade_inputs(payload)
    if not student_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No student text could be extracted from the uploaded file.",
        )
    custom_prompt = str(payload.custom_prompt or "").strip()
    prompt = USER_PROMPT_TEMPLATE.format(
        guide=guide_text,
        student_answers=student_text,
    )
    if custom_prompt:
        prompt = (
            prompt
            + "\n\nNOTE: The user provided an additional custom prompt. "
            "Follow it as long as it does not conflict with the required JSON format and schema.\n"
            f"CUSTOM PROMPT:\n{custom_prompt}"
        )

    try:
        agent = get_agent()
        logfire.info(
            "Grading request received",
            guide_chars=len(guide_text),
            student_chars=len(student_text),
            bypass_validation=BYPASS_OUTPUT_VALIDATION,
        )
        LOCAL = False
        if BYPASS_OUTPUT_VALIDATION:
            if LOCAL:
                result = local_ai_nemo(prompt=prompt)
            
            else:
                result = agent.run_sync(prompt, output_type=str)
            
            
            if isinstance(result, str):
                raw_output = result
            elif isinstance(getattr(result, "output", None), str):
                raw_output = result.output
            else:
                raw_output = _extract_raw_output(result)
            raw_output_clean = _strip_code_fences(raw_output)
            _attach_raw_output(raw_output_clean)
            run_id = _store_raw_output(raw_output_clean)
            response.headers["X-Grady-Run-Id"] = run_id

            grading_view = _extract_grading_view(raw_output_clean)
            if grading_view is None:
                logfire.warning("Failed to parse grading view; returning raw output")
                print("\nRaw model output (pre-validation, cleaned):", raw_output_clean)
                print("1:", result)
                return Response(content=raw_output_clean, media_type="application/json")
            _attach_questions_map(grading_view.questions)
            _attach_answers_map(grading_view.answers)
            _attach_student_name(grading_view.student_name)
            _attach_grades_map(grading_view.grades)
            _attach_comments_map(grading_view.comments)
            _attach_highlights_map(grading_view.highlights)
            _attach_criteria_list(grading_view.criteria)
            logfire.info("Raw model output (pre-validation)", raw_output=raw_output_clean)
            print("\nRaw model output (pre-validation, cleaned):", raw_output_clean)
            print("1:", result)
            return grading_view
        
        
        LOCAL = False
        if LOCAL:
            raw_output = local_ai_nemo(prompt)
            raw_output_clean = _strip_code_fences(raw_output)
            _attach_raw_output(raw_output_clean)
            run_id = _store_raw_output(raw_output_clean)
            response.headers["X-Grady-Run-Id"] = run_id

            grading_view = _extract_grading_view(raw_output_clean)
            if grading_view is None:
                logfire.warning("Failed to parse grading view; returning raw output")
                print("\nRaw model output (pre-validation, cleaned):", raw_output_clean)
                return Response(content=raw_output_clean, media_type="application/json")
            _attach_questions_map(grading_view.questions)
            _attach_answers_map(grading_view.answers)
            _attach_student_name(grading_view.student_name)
            _attach_grades_map(grading_view.grades)
            _attach_comments_map(grading_view.comments)
            _attach_highlights_map(grading_view.highlights)
            _attach_criteria_list(grading_view.criteria)
            logfire.info("Raw model output (pre-validation)", raw_output=raw_output_clean)
            print("\nRaw model output (pre-validation, cleaned):", raw_output_clean)
            print("1:", raw_output_clean)
            return grading_view

        result = agent.run_sync(prompt)
        raw_output = _extract_raw_output(result)
        raw_output_clean = _strip_code_fences(raw_output)
        _attach_raw_output(raw_output_clean)
        _attach_questions_map(result.output.questions)
        _attach_answers_map(result.output.answers)
        _attach_student_name(result.output.student_name)
        _attach_grades_map(result.output.grades)
        _attach_comments_map(result.output.comments)
        _attach_highlights_map(result.output.highlights)
        _attach_criteria_list(result.output.criteria)
        logfire.info("Raw model output (pre-validation)", raw_output=raw_output_clean)
        print("\nRaw model output (pre-validation, cleaned):", raw_output_clean)
        print("1:", result)
        print("\nGrades result:", result.output)
        run_id = _store_raw_output(raw_output_clean)
        response.headers["X-Grady-Run-Id"] = run_id
        return result.output
    except Exception as exc:
        logfire.exception("Grading failed")
        raise HTTPException(status_code=500, detail=f"Grading failed: {exc}") from exc


@app.post("/grade/stream")
async def grade_stream(payload: GradeRequest) -> StreamingResponse:
    guide_text, student_text = _resolve_grade_inputs(payload)
    if not student_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No student text could be extracted from the uploaded file.",
        )
    custom_prompt = str(payload.custom_prompt or "").strip()
    prompt = USER_PROMPT_TEMPLATE.format(
        guide=guide_text,
        student_answers=student_text,
    )
    if custom_prompt:
        prompt = (
            prompt
            + "\n\nNOTE: The user provided an additional custom prompt. "
            "Follow it as long as it does not conflict with the required JSON format and schema.\n"
            f"CUSTOM PROMPT:\n{custom_prompt}"
        )

    async def event_stream():
        try:
            agent = get_agent()
            logfire.info(
                "Grading stream request received",
                guide_chars=len(guide_text),
                student_chars=len(student_text),
                bypass_validation=True,
            )
            yield _ndjson_event({"type": "status", "stage": "start"})

            async with agent.run_stream(prompt, output_type=str) as stream:
                async for chunk in stream.stream_text(delta=True):
                    if chunk:
                        yield _ndjson_event({"type": "text", "delta": chunk})
                raw_output = await stream.get_output()

            raw_output_clean = _strip_code_fences(raw_output)
            _attach_raw_output(raw_output_clean)
            run_id = _store_raw_output(raw_output_clean)

            grading_view = _extract_grading_view(raw_output_clean)
            if grading_view is None:
                logfire.warning("Failed to parse grading view; streaming raw output only")
                print("\nRaw model output (stream, cleaned):", raw_output_clean)
                yield _ndjson_event(
                    {
                        "type": "raw",
                        "run_id": run_id,
                        "raw_output": raw_output_clean,
                    }
                )
                return

            _attach_questions_map(grading_view.questions)
            _attach_answers_map(grading_view.answers)
            _attach_student_name(grading_view.student_name)
            _attach_grades_map(grading_view.grades)
            _attach_comments_map(grading_view.comments)
            _attach_highlights_map(grading_view.highlights)
            _attach_criteria_list(grading_view.criteria)
            logfire.info("Raw model output (stream, cleaned)", raw_output=raw_output_clean)
            print("\nRaw model output (stream, cleaned):", raw_output_clean)
            yield _ndjson_event(
                {
                    "type": "final",
                    "run_id": run_id,
                    "raw_output": raw_output_clean,
                    "data": grading_view.model_dump(),
                }
            )
        except Exception as exc:
            logfire.exception("Streaming grading failed")
            yield _ndjson_event({"type": "error", "detail": f"Grading failed: {exc}"})

    return StreamingResponse(
        event_stream(),
        media_type="application/x-ndjson",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/{path:path}")
def frontend_fallback(path: str, request: Request) -> Response:
    return _serve_frontend(request, path)
